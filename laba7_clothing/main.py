from kivy.app import App
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.label import Label
from kivy.uix.button import Button
from kivy.uix.spinner import Spinner
from kivy.uix.textinput import TextInput
from kivy.core.window import Window
from docx import Document
from openpyxl import Workbook
import json

from clothing_package import CoatCalculator, TrousersCalculator, SuitCalculator

Window.size = (550, 650)

class ClothingApp(App):
    def build(self):
        return MainWidget()

class MainWidget(BoxLayout):
    def __init__(self, **kwargs):
        super().__init__(**kwargs)
        self.orientation = 'vertical'
        self.padding = 10
        self.spacing = 10
        self.current_result = None
        self.current_calculator = None
        self._build_ui()
    
    def _build_ui(self):
        заголовок = Label(text="Расчёт стоимости пошива одежды\nВариант 10: Одежда", font_size='18sp', size_hint=(1, 0.1))
        self.add_widget(заголовок)
        
        тип_ряда = BoxLayout(size_hint=(1, 0.08))
        тип_надпись = Label(text="Тип одежды:", size_hint=(0.3, 1), halign='right')
        тип_ряда.add_widget(тип_надпись)
        self.выбор_типа = Spinner(text="Пиджак", values=("Пиджак", "Брюки", "Костюм-тройка"), size_hint=(0.7, 1))
        тип_ряда.add_widget(self.выбор_типа)
        self.add_widget(тип_ряда)
        
        размер_ряда = BoxLayout(size_hint=(1, 0.08))
        размер_надпись = Label(text="Размер (44-54):", size_hint=(0.3, 1), halign='right')
        размер_ряда.add_widget(размер_надпись)
        self.поле_размера = TextInput(text="44", input_filter='int', multiline=False, size_hint=(0.7, 1))
        размер_ряда.add_widget(self.поле_размера)
        self.add_widget(размер_ряда)
        
        self.кнопка_расчёта = Button(text="РАССЧИТАТЬ", size_hint=(1, 0.08), background_color=(0.4, 0.4, 0.8, 1))
        self.кнопка_расчёта.bind(on_press=self.рассчитать)
        self.add_widget(self.кнопка_расчёта)
        
        self.поле_результата = TextInput(readonly=True, multiline=True, size_hint=(1, 0.5))
        self.поле_результата.background_color = (0.95, 0.95, 0.95, 1)
        self.add_widget(self.поле_результата)
        
        кнопки_ряда = BoxLayout(size_hint=(1, 0.08), spacing=10)
        self.кнопка_doc = Button(text="Сохранить в DOC", disabled=True, background_color=(0.5, 0.2, 0.2, 1))
        self.кнопка_doc.bind(on_press=self.сохранить_doc)
        кнопки_ряда.add_widget(self.кнопка_doc)
        self.кнопка_xls = Button(text="Сохранить в XLS", disabled=True, background_color=(0.2, 0.4, 0.6, 1))
        self.кнопка_xls.bind(on_press=self.сохранить_xls)
        кнопки_ряда.add_widget(self.кнопка_xls)
        self.add_widget(кнопки_ряда)
    
    def показать_результат(self, результат):
        """Красивое форматирование результата без боковых стенок"""
        if 'coat' in результат:
            return f"""
==========================================
    РЕЗУЛЬТАТ РАСЧЁТА - КОСТЮМ-ТРОЙКА        
==========================================

  Размер: {результат['size']}                                              
__________________________________________

  ПИДЖАК:                                      
    Ткань: {результат['coat']['fabric_meters']} м | Стоимость: {результат['coat']['total']} руб    
  БРЮКИ:                                       
    Ткань: {результат['trousers']['fabric_meters']} м | Стоимость: {результат['trousers']['total']} руб  
  ЖИЛЕТ:                                       
    Ткань: {результат['vest']['fabric_meters']} м | Стоимость: {результат['vest']['total']} руб    
__________________________________________

  ОБЩАЯ СТОИМОСТЬ: {результат['total']} руб                      
==========================================
"""
        else:
            return f"""
==========================================
            РЕЗУЛЬТАТ РАСЧЁТА                
==========================================

  Тип одежды:     {результат['type']:<25}
  Размер:         {результат['size']}                                        
__________________________________________

  Расход ткани:   {результат['fabric_meters']} м                          
  Стоимость ткани: {результат['fabric_cost']:>8} руб                       
  Стоимость работы: {результат['work_cost']:>7} руб                       
  Фурнитура:      {результат['accessories']:>8} руб                       
__________________________________________

  ИТОГО:          {результат['total']:>10} руб                     
==========================================
"""
    
    def рассчитать(self, instance):
        try:
            размер = int(self.поле_размера.text)
            if размер < 44 or размер > 54:
                self.показать_ошибку("Размер должен быть от 44 до 54")
                return
            тип_одежды = self.выбор_типа.text
            if тип_одежды == "Пиджак":
                self.текущий_калькулятор = CoatCalculator(размер)
            elif тип_одежды == "Брюки":
                self.текущий_калькулятор = TrousersCalculator(размер)
            else:
                self.текущий_калькулятор = SuitCalculator(размер)
            self.текущий_результат = self.текущий_калькулятор.get_result_dict()
            
            self.поле_результата.text = self.показать_результат(self.текущий_результат)
            
            self.кнопка_doc.disabled = False
            self.кнопка_xls.disabled = False
        except Exception as e:
            self.показать_ошибку(f"Ошибка: {e}")
    
    def сохранить_doc(self, instance):
        if not self.текущий_результат:
            self.показать_ошибку("Сначала выполните расчёт!")
            return
        try:
            тип = self.текущий_результат['type']
            размер = self.текущий_результат['size']
            имя_файла = f"{тип}_{размер}.docx"
            self.создать_doc_файл(имя_файла)
            self.показать_сообщение(f"Сохранено в {имя_файла}")
        except Exception as e:
            self.показать_ошибку(f"Ошибка сохранения DOC: {e}")
    
    def сохранить_xls(self, instance):
        if not self.текущий_результат:
            self.показать_ошибку("Сначала выполните расчёт!")
            return
        try:
            тип = self.текущий_результат['type']
            размер = self.текущий_результат['size']
            имя_файла = f"{тип}_{размер}.xlsx"
            self.создать_xls_файл(имя_файла)
            self.показать_сообщение(f"Сохранено в {имя_файла}")
        except Exception as e:
            self.показать_ошибку(f"Ошибка сохранения XLS: {e}")
    
    def создать_doc_файл(self, имя_файла):
        doc = Document()
        doc.add_heading('Расчёт стоимости пошива одежды', 0)
        doc.add_heading('Параметры заказа', level=1)
        doc.add_paragraph(f"Тип одежды: {self.текущий_результат.get('type', 'Костюм-тройка')}")
        doc.add_paragraph(f"Размер: {self.текущий_результат['size']}")
        doc.add_heading('Детальный расчёт', level=1)
        
        if 'coat' in self.текущий_результат:
            for часть in ['coat', 'trousers', 'vest']:
                название = "Пиджак" if часть == "coat" else "Брюки" if часть == "trousers" else "Жилет"
                doc.add_heading(название, level=2)
                doc.add_paragraph(f"Расход ткани: {self.текущий_результат[часть]['fabric_meters']} м")
                doc.add_paragraph(f"Стоимость ткани: {self.текущий_результат[часть]['fabric_cost']} руб")
                doc.add_paragraph(f"Стоимость работы: {self.текущий_результат[часть]['work_cost']} руб")
                doc.add_paragraph(f"Фурнитура: {self.текущий_результат[часть]['accessories']} руб")
                doc.add_paragraph(f"Итого {название}: {self.текущий_результат[часть]['total']} руб")
        else:
            doc.add_paragraph(f"Расход ткани: {self.текущий_результат['fabric_meters']} м")
            doc.add_paragraph(f"Стоимость ткани: {self.текущий_результат['fabric_cost']} руб")
            doc.add_paragraph(f"Стоимость работы: {self.текущий_результат['work_cost']} руб")
            doc.add_paragraph(f"Фурнитура: {self.текущий_результат['accessories']} руб")
        
        doc.add_heading('Общая стоимость', level=1)
        doc.add_paragraph(f"ИТОГО: {self.текущий_результат['total']} рублей")
        doc.save(имя_файла)
    
    def создать_xls_файл(self, имя_файла):
        wb = Workbook()
        ws = wb.active
        ws.title = "Результат"
        
        if 'coat' in self.текущий_результат:
            ws['A1'] = "Расчёт стоимости пошива костюма-тройки"
            ws['A3'] = "Параметр"
            ws['B3'] = "Пиджак"
            ws['C3'] = "Брюки"
            ws['D3'] = "Жилет"
            
            ws['A4'] = "Размер"
            ws['B4'] = self.текущий_результат['size']
            ws['C4'] = self.текущий_результат['size']
            ws['D4'] = self.текущий_результат['size']
            
            ws['A5'] = "Расход ткани (м)"
            ws['B5'] = self.текущий_результат['coat']['fabric_meters']
            ws['C5'] = self.текущий_результат['trousers']['fabric_meters']
            ws['D5'] = self.текущий_результат['vest']['fabric_meters']
            
            ws['A6'] = "Стоимость ткани (руб)"
            ws['B6'] = self.текущий_результат['coat']['fabric_cost']
            ws['C6'] = self.текущий_результат['trousers']['fabric_cost']
            ws['D6'] = self.текущий_результат['vest']['fabric_cost']
            
            ws['A7'] = "Стоимость работы (руб)"
            ws['B7'] = self.текущий_результат['coat']['work_cost']
            ws['C7'] = self.текущий_результат['trousers']['work_cost']
            ws['D7'] = self.текущий_результат['vest']['work_cost']
            
            ws['A8'] = "Фурнитура (руб)"
            ws['B8'] = self.текущий_результат['coat']['accessories']
            ws['C8'] = self.текущий_результат['trousers']['accessories']
            ws['D8'] = self.текущий_результат['vest']['accessories']
            
            ws['A9'] = "Итого (руб)"
            ws['B9'] = self.текущий_результат['coat']['total']
            ws['C9'] = self.текущий_результат['trousers']['total']
            ws['D9'] = self.текущий_результат['vest']['total']
            
            ws['A11'] = "ОБЩАЯ СТОИМОСТЬ:"
            ws['B11'] = self.текущий_результат['total']
        else:
            ws['A1'] = f"Расчёт стоимости пошива {self.текущий_результат['type']}"
            данные = [
                ["Тип одежды", self.текущий_результат['type']],
                ["Размер", self.текущий_результат['size']],
                ["Расход ткани (м)", self.текущий_результат['fabric_meters']],
                ["Стоимость ткани (руб)", self.текущий_результат['fabric_cost']],
                ["Стоимость работы (руб)", self.текущий_результат['work_cost']],
                ["Фурнитура (руб)", self.текущий_результат['accessories']],
                ["ИТОГО (руб)", self.текущий_результат['total']]
            ]
            for i, (параметр, значение) in enumerate(данные, start=3):
                ws.cell(row=i, column=1, value=параметр)
                ws.cell(row=i, column=2, value=значение)
        
        wb.save(имя_файла)
    
    def показать_ошибку(self, сообщение):
        from kivy.uix.popup import Popup
        всплывающее_окно = Popup(title='Ошибка', content=Label(text=сообщение), size_hint=(0.8, 0.3))
        всплывающее_окно.open()
    
    def показать_сообщение(self, сообщение):
        from kivy.uix.popup import Popup
        всплывающее_окно = Popup(title='Успех', content=Label(text=сообщение), size_hint=(0.8, 0.3))
        всплывающее_окно.open()

if __name__ == '__main__':
    ClothingApp().run()