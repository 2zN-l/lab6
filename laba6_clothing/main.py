import tkinter as tk
from tkinter import ttk, messagebox
from clothing_package import CoatCalculator, TrousersCalculator, SuitCalculator
from docx import Document
from openpyxl import Workbook

class ClothingApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Расчёт стоимости пошива одежды")
        self.root.geometry("550x650")
        self.root.resizable(False, False)
        self.clothing_type = tk.StringVar(value="Пиджак")
        self.size = tk.IntVar(value=44)
        self.setup_ui()
    
    def setup_ui(self):
        title = tk.Label(self.root, text="Расчёт расхода ткани и стоимости пошива", font=("Arial", 16, "bold"), fg="#000000")
        title.pack(pady=15)
        
        input_frame = tk.LabelFrame(self.root, text="Данные заказа", font=("Arial", 12), padx=20, pady=15)
        input_frame.pack(fill="x", padx=20, pady=10)
        
        tk.Label(input_frame, text="Тип одежды:", font=("Arial", 11)).grid(row=0, column=0, sticky="w", pady=5)
        clothing_options = ["Пиджак", "Брюки", "Костюм-тройка"]
        clothing_menu = ttk.Combobox(input_frame, textvariable=self.clothing_type, values=clothing_options, state="readonly", width=20)
        clothing_menu.grid(row=0, column=1, pady=5, padx=10)
        
        tk.Label(input_frame, text="Размер (44-54):", font=("Arial", 11)).grid(row=1, column=0, sticky="w", pady=5)
        size_spinbox = tk.Spinbox(input_frame, from_=44, to=54, textvariable=self.size, width=20, font=("Arial", 11))
        size_spinbox.grid(row=1, column=1, pady=5, padx=10)
        
        self.calc_button = tk.Button(input_frame, text="РАССЧИТАТЬ", command=self.calculate, bg="#000000", fg="white", font=("Arial", 8, "bold"), padx=5, pady=1)
        self.calc_button.grid(row=2, column=0, columnspan=1, pady=2, sticky="w")
        
        self.result_frame = tk.LabelFrame(self.root, text="Результат", font=("Arial", 12), padx=15, pady=10)
        self.result_frame.pack(fill="both", expand=True, padx=20, pady=10)
        
        self.result_text = tk.Text(self.result_frame, height=12, font=("Courier", 10))
        self.result_text.pack(fill="both", expand=True)
        
        scrollbar = tk.Scrollbar(self.result_text)
        scrollbar.pack(side="right", fill="y")
        self.result_text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.result_text.yview)
        
        save_frame = tk.Frame(self.root)
        save_frame.pack(fill="x", padx=20, pady=10)
        
        self.save_doc_button = tk.Button(save_frame, text="Сохранить в DOC", command=self.save_to_doc, bg="#42001f", fg="white", font=("Arial", 11), padx=15, pady=5, state="disabled")
        self.save_doc_button.pack(side="left", padx=5)
        
        self.save_xls_button = tk.Button(save_frame, text="Сохранить в XLS", command=self.save_to_xls, bg="#3b4e5e", fg="white", font=("Arial", 11), padx=15, pady=5, state="disabled")
        self.save_xls_button.pack(side="left", padx=5)
        
        self.current_result = None
    
    def calculate(self):
        size = self.size.get()
        clothing = self.clothing_type.get()
        
        try:
            if clothing == "Пиджак":
                self.current_result = CoatCalculator.calculate_price(size)
                self.display_result(self.current_result)
            elif clothing == "Брюки":
                self.current_result = TrousersCalculator.calculate_price(size)
                self.display_result(self.current_result)
            elif clothing == "Костюм-тройка":
                self.current_result = SuitCalculator.calculate_price(size)
                self.display_suit_result(self.current_result)
            
            self.save_doc_button.config(state="normal")
            self.save_xls_button.config(state="normal")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка расчёта: {str(e)}")
    
    def display_result(self, result):
        self.result_text.delete(1.0, tk.END)
        text = f"""
==========================================
            РЕЗУЛЬТАТ РАСЧЁТА                
==========================================

  Тип одежды:     {result['type']:<25}
  Размер:         {result['size']}                                        
__________________________________________

  Расход ткани:   {result['fabric_meters']} м                          
  Стоимость ткани:{result['fabric_cost']:>8} руб                       
  Стоимость работы:{result['work_cost']:>7} руб                       
  Фурнитура:      {result['accessories']:>8} руб                       
__________________________________________

  ИТОГО:          {result['total']:>10} руб                     
==========================================
        """
        self.result_text.insert(1.0, text)
    
    def display_suit_result(self, result):
        self.result_text.delete(1.0, tk.END)
        text = f"""
==========================================
    РЕЗУЛЬТАТ РАСЧЁТА - КОСТЮМ-ТРОЙКА        
==========================================

  Размер: {result['size']}                                              
__________________________________________

  ПИДЖАК:                                      
    Ткань: {result['coat']['fabric_meters']} м | Стоимость: {result['coat']['total']} руб    
  БРЮКИ:                                       
    Ткань: {result['trousers']['fabric_meters']} м | Стоимость: {result['trousers']['total']} руб  
  ЖИЛЕТ:                                       
    Ткань: {result['vest']['fabric_meters']} м | Стоимость: {result['vest']['total']} руб    
__________________________________________

 ОБЩАЯ СТОИМОСТЬ: {result['total']} руб                      
==========================================
        """
        self.result_text.insert(1.0, text)
    
    def save_to_doc(self):
        if not self.current_result:
            messagebox.showwarning("Предупреждение", "Сначала выполните расчёт!")
            return
        
        try:
            doc = Document()
            doc.add_heading('Расчёт стоимости пошива одежды', 0)
            
            doc.add_heading('Параметры заказа', level=1)
            doc.add_paragraph(f"Тип одежды: {self.current_result.get('type', 'Костюм-тройка')}")
            doc.add_paragraph(f"Размер: {self.current_result['size']}")
            
            doc.add_heading('Детальный расчёт', level=1)
            
            if 'coat' in self.current_result:
                doc.add_heading('Пиджак', level=2)
                doc.add_paragraph(f"Расход ткани: {self.current_result['coat']['fabric_meters']} м")
                doc.add_paragraph(f"Стоимость ткани: {self.current_result['coat']['fabric_cost']} руб")
                doc.add_paragraph(f"Стоимость работы: {self.current_result['coat']['work_cost']} руб")
                doc.add_paragraph(f"Фурнитура: {self.current_result['coat']['accessories']} руб")
                doc.add_paragraph(f"Итого пиджак: {self.current_result['coat']['total']} руб")
                
                doc.add_heading('Брюки', level=2)
                doc.add_paragraph(f"Расход ткани: {self.current_result['trousers']['fabric_meters']} м")
                doc.add_paragraph(f"Стоимость ткани: {self.current_result['trousers']['fabric_cost']} руб")
                doc.add_paragraph(f"Стоимость работы: {self.current_result['trousers']['work_cost']} руб")
                doc.add_paragraph(f"Фурнитура: {self.current_result['trousers']['accessories']} руб")
                doc.add_paragraph(f"Итого брюки: {self.current_result['trousers']['total']} руб")
                
                doc.add_heading('Жилет', level=2)
                doc.add_paragraph(f"Расход ткани: {self.current_result['vest']['fabric_meters']} м")
                doc.add_paragraph(f"Стоимость ткани: {self.current_result['vest']['fabric_cost']} руб")
                doc.add_paragraph(f"Стоимость работы: {self.current_result['vest']['work_cost']} руб")
                doc.add_paragraph(f"Фурнитура: {self.current_result['vest']['accessories']} руб")
                doc.add_paragraph(f"Итого жилет: {self.current_result['vest']['total']} руб")
                
                doc.add_heading('Общая стоимость', level=1)
                doc.add_paragraph(f"ИТОГО: {self.current_result['total']} рублей")
            else:
                doc.add_paragraph(f"Расход ткани: {self.current_result['fabric_meters']} м")
                doc.add_paragraph(f"Стоимость ткани: {self.current_result['fabric_cost']} руб")
                doc.add_paragraph(f"Стоимость работы: {self.current_result['work_cost']} руб")
                doc.add_paragraph(f"Фурнитура: {self.current_result['accessories']} руб")
                doc.add_paragraph(f"ИТОГО: {self.current_result['total']} руб")
            
            filename = f"report_{self.current_result['type']}_{self.current_result['size']}.docx"
            doc.save(filename)
            messagebox.showinfo("Успех", f"Отчёт сохранён в файл:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить DOC:\n{str(e)}")
    
    def save_to_xls(self):
        if not self.current_result:
            messagebox.showwarning("Предупреждение", "Сначала выполните расчёт!")
            return
        
        try:
            wb = Workbook()
            
            if 'coat' in self.current_result:
                ws = wb.active
                ws.title = "Костюм-тройка"
                
                ws['A1'] = "Расчёт стоимости пошива костюма-тройки"
                
                ws['A3'] = "Параметр"
                ws['B3'] = "Пиджак"
                ws['C3'] = "Брюки"
                ws['D3'] = "Жилет"
                
                ws['A4'] = "Размер"
                ws['B4'] = self.current_result['size']
                ws['C4'] = self.current_result['size']
                ws['D4'] = self.current_result['size']
                
                ws['A5'] = "Расход ткани (м)"
                ws['B5'] = self.current_result['coat']['fabric_meters']
                ws['C5'] = self.current_result['trousers']['fabric_meters']
                ws['D5'] = self.current_result['vest']['fabric_meters']
                
                ws['A6'] = "Стоимость ткани (руб)"
                ws['B6'] = self.current_result['coat']['fabric_cost']
                ws['C6'] = self.current_result['trousers']['fabric_cost']
                ws['D6'] = self.current_result['vest']['fabric_cost']
                
                ws['A7'] = "Стоимость работы (руб)"
                ws['B7'] = self.current_result['coat']['work_cost']
                ws['C7'] = self.current_result['trousers']['work_cost']
                ws['D7'] = self.current_result['vest']['work_cost']
                
                ws['A8'] = "Фурнитура (руб)"
                ws['B8'] = self.current_result['coat']['accessories']
                ws['C8'] = self.current_result['trousers']['accessories']
                ws['D8'] = self.current_result['vest']['accessories']
                
                ws['A9'] = "Итого (руб)"
                ws['B9'] = self.current_result['coat']['total']
                ws['C9'] = self.current_result['trousers']['total']
                ws['D9'] = self.current_result['vest']['total']
                
                ws['A11'] = "ОБЩАЯ СТОИМОСТЬ:"
                ws['B11'] = self.current_result['total']
                
            else:
                ws = wb.active
                ws.title = self.current_result['type']
                
                ws['A1'] = f"Расчёт стоимости пошива {self.current_result['type']}"
                
                ws['A3'] = "Параметр"
                ws['B3'] = "Значение"
                
                data = [
                    ["Тип одежды", self.current_result['type']],
                    ["Размер", self.current_result['size']],
                    ["Расход ткани (м)", self.current_result['fabric_meters']],
                    ["Стоимость ткани (руб)", self.current_result['fabric_cost']],
                    ["Стоимость работы (руб)", self.current_result['work_cost']],
                    ["Фурнитура (руб)", self.current_result['accessories']],
                    ["ИТОГО (руб)", self.current_result['total']]
                ]
                
                for i, (param, value) in enumerate(data, start=4):
                    ws[f'A{i}'] = param
                    ws[f'B{i}'] = value
            
            filename = f"report_{self.current_result['type']}_{self.current_result['size']}.xlsx"
            wb.save(filename)
            messagebox.showinfo("Успех", f"Отчёт сохранён в файл:\n{filename}")
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось сохранить XLS:\n{str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = ClothingApp(root)
    root.mainloop()