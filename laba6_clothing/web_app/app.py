from flask import Flask, request, render_template, flash, redirect, url_for, send_file, session
from clothing_package import *
from clothing_package.db import DatabaseManager
import json
import io
import os
import sys

sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

app = Flask(__name__)
app.secret_key = 'your-secret-key-here-change-in-production'
db_manager = DatabaseManager()

def format_result_for_display(result):
    if 'coat' in result:
        return f"""
==========================================
    РЕЗУЛЬТАТ РАСЧЁТА - КОСТЮМ-ТРОЙКА        
==========================================

  Размер: {result['size']}                                              
__________________________________________

  ПИДЖАК:                                      
    Ткань: {result['coat']['fabric_meters']} м | Стоимость: {result['coat']['total']} руб    
  БРЮКИ:                                       
    Ткань: {result['trousers']['fabric_meters']} м | Стоимость: {result['trousers']['total']} руб  
  ЖИЛЕТ:                                       
    Ткань: {result['vest']['fabric_meters']} м | Стоимость: {result['vest']['total']} руб    
__________________________________________

  ОБЩАЯ СТОИМОСТЬ: {result['total']} руб                      
==========================================
        """
    else:
        return f"""
==========================================
            РЕЗУЛЬТАТ РАСЧЁТА                
==========================================

  Тип одежды:     {result['type']:<25}
  Размер:         {result['size']}                                        
__________________________________________

  Расход ткани:   {result['fabric_meters']} м                          
  Стоимость ткани:{result['fabric_cost']:>8} руб                       
  Стоимость работы:{result['work_cost']:>7} руб                       
  Фурнитура:      {result['accessories']:>8} руб                       
__________________________________________

  ИТОГО:          {result['total']:>10} руб                     
==========================================
        """

def create_doc_file(result):
    from docx import Document
    doc = Document()
    doc.add_heading('Расчёт стоимости пошива одежды', 0)
    doc.add_heading('Параметры заказа', level=1)
    doc.add_paragraph(f"Тип одежды: {result.get('type', 'Костюм-тройка')}")
    doc.add_paragraph(f"Размер: {result['size']}")
    doc.add_heading('Детальный расчёт', level=1)
    
    if 'coat' in result:
        doc.add_heading('Пиджак', level=2)
        doc.add_paragraph(f"Расход ткани: {result['coat']['fabric_meters']} м")
        doc.add_paragraph(f"Стоимость ткани: {result['coat']['fabric_cost']} руб")
        doc.add_paragraph(f"Стоимость работы: {result['coat']['work_cost']} руб")
        doc.add_paragraph(f"Фурнитура: {result['coat']['accessories']} руб")
        doc.add_paragraph(f"Итого пиджак: {result['coat']['total']} руб")
        doc.add_heading('Брюки', level=2)
        doc.add_paragraph(f"Расход ткани: {result['trousers']['fabric_meters']} м")
        doc.add_paragraph(f"Стоимость ткани: {result['trousers']['fabric_cost']} руб")
        doc.add_paragraph(f"Стоимость работы: {result['trousers']['work_cost']} руб")
        doc.add_paragraph(f"Фурнитура: {result['trousers']['accessories']} руб")
        doc.add_paragraph(f"Итого брюки: {result['trousers']['total']} руб")
        doc.add_heading('Жилет', level=2)
        doc.add_paragraph(f"Расход ткани: {result['vest']['fabric_meters']} м")
        doc.add_paragraph(f"Стоимость ткани: {result['vest']['fabric_cost']} руб")
        doc.add_paragraph(f"Стоимость работы: {result['vest']['work_cost']} руб")
        doc.add_paragraph(f"Фурнитура: {result['vest']['accessories']} руб")
        doc.add_paragraph(f"Итого жилет: {result['vest']['total']} руб")
        doc.add_heading('Общая стоимость', level=1)
        doc.add_paragraph(f"ИТОГО: {result['total']} рублей")
    else:
        doc.add_paragraph(f"Расход ткани: {result['fabric_meters']} м")
        doc.add_paragraph(f"Стоимость ткани: {result['fabric_cost']} руб")
        doc.add_paragraph(f"Стоимость работы: {result['work_cost']} руб")
        doc.add_paragraph(f"Фурнитура: {result['accessories']} руб")
        doc.add_paragraph(f"ИТОГО: {result['total']} руб")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def create_xls_file(result):
    from openpyxl import Workbook
    wb = Workbook()
    
    if 'coat' in result:
        ws = wb.active
        ws.title = "Костюм-тройка"
        ws['A1'] = "Расчёт стоимости пошива костюма-тройки"
        ws['A3'] = "Параметр"
        ws['B3'] = "Пиджак"
        ws['C3'] = "Брюки"
        ws['D3'] = "Жилет"
        ws['A4'] = "Размер"
        ws['B4'] = result['size']
        ws['C4'] = result['size']
        ws['D4'] = result['size']
        ws['A5'] = "Расход ткани (м)"
        ws['B5'] = result['coat']['fabric_meters']
        ws['C5'] = result['trousers']['fabric_meters']
        ws['D5'] = result['vest']['fabric_meters']
        ws['A6'] = "Стоимость ткани (руб)"
        ws['B6'] = result['coat']['fabric_cost']
        ws['C6'] = result['trousers']['fabric_cost']
        ws['D6'] = result['vest']['fabric_cost']
        ws['A7'] = "Стоимость работы (руб)"
        ws['B7'] = result['coat']['work_cost']
        ws['C7'] = result['trousers']['work_cost']
        ws['D7'] = result['vest']['work_cost']
        ws['A8'] = "Фурнитура (руб)"
        ws['B8'] = result['coat']['accessories']
        ws['C8'] = result['trousers']['accessories']
        ws['D8'] = result['vest']['accessories']
        ws['A9'] = "Итого (руб)"
        ws['B9'] = result['coat']['total']
        ws['C9'] = result['trousers']['total']
        ws['D9'] = result['vest']['total']
        ws['A11'] = "ОБЩАЯ СТОИМОСТЬ:"
        ws['B11'] = result['total']
    else:
        ws = wb.active
        ws.title = result['type']
        ws['A1'] = f"Расчёт стоимости пошива {result['type']}"
        ws['A3'] = "Параметр"
        ws['B3'] = "Значение"
        data = [
            ["Тип одежды", result['type']],
            ["Размер", result['size']],
            ["Расход ткани (м)", result['fabric_meters']],
            ["Стоимость ткани (руб)", result['fabric_cost']],
            ["Стоимость работы (руб)", result['work_cost']],
            ["Фурнитура (руб)", result['accessories']],
            ["ИТОГО (руб)", result['total']]
        ]
        for i, (param, value) in enumerate(data, start=4):
            ws[f'A{i}'] = param
            ws[f'B{i}'] = value
    
    file_stream = io.BytesIO()
    wb.save(file_stream)
    file_stream.seek(0)
    return file_stream

@app.route('/', methods=['GET', 'POST'])
def index():
    result = None
    result_text = None
    result_json = None
    clothing_type = 'Пиджак'
    size = 44
    
    if request.method == 'POST':
        clothing_type = request.form.get('clothing_type')
        size = int(request.form.get('size', 44))
        
        try:
            if clothing_type == 'Пиджак':
                calculator = CoatCalculator()
                result = calculator.calculate_price(size)
            elif clothing_type == 'Брюки':
                calculator = TrousersCalculator()
                result = calculator.calculate_price(size)
            else:
                calculator = SuitCalculator()
                result = calculator.calculate_price(size)
            
            result_text = format_result_for_display(result)
            result_json = json.dumps(result, ensure_ascii=False)
            session['last_result'] = result_json
            flash('Расчёт выполнен успешно!', 'success')
            
        except Exception as e:
            flash(f'Ошибка расчёта: {str(e)}', 'error')
    
    return render_template('index.html', result=result, result_text=result_text, result_json=result_json, clothing_type=clothing_type, size=size)

@app.route('/save_to_db', methods=['POST'])
def save_to_db():
    result_json = request.form.get('result_json')
    if not result_json:
        flash('Нет данных для сохранения', 'error')
        return redirect(url_for('index'))
    
    try:
        result = json.loads(result_json)
        success = db_manager.save_result(result)
        if success:
            flash('Результат успешно сохранён в базу данных PostgreSQL!', 'success')
        else:
            flash('Ошибка при сохранении в БД. Проверьте, запущен ли контейнер Docker.', 'error')
    except Exception as e:
        flash(f'Ошибка при сохранении в БД: {str(e)}', 'error')
    
    return redirect(url_for('index'))

@app.route('/export_doc')
def export_doc():
    result_json = request.args.get('result_json')
    if not result_json:
        flash('Нет данных для экспорта', 'error')
        return redirect(url_for('index'))
    
    try:
        result = json.loads(result_json)
        file_stream = create_doc_file(result)
        return send_file(file_stream, as_attachment=True, download_name=f"report_{result['type']}_{result['size']}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        flash(f'Ошибка при экспорте в DOC: {str(e)}', 'error')
        return redirect(url_for('index'))

@app.route('/export_xls')
def export_xls():
    result_json = request.args.get('result_json')
    if not result_json:
        flash('Нет данных для экспорта', 'error')
        return redirect(url_for('index'))
    
    try:
        result = json.loads(result_json)
        file_stream = create_xls_file(result)
        return send_file(file_stream, as_attachment=True, download_name=f"report_{result['type']}_{result['size']}.xlsx", mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    except Exception as e:
        flash(f'Ошибка при экспорте в XLS: {str(e)}', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True, port=5000)